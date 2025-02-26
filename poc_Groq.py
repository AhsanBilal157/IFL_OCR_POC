import os
import streamlit as st
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain.chains.llm import LLMChain
from langchain_groq import ChatGroq
import json
import re


# Load environment variables
load_dotenv()

# Azure API configuration
AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT = os.getenv("DOCUMENT_INTELLIGENCE_ENDPOINT")
AZURE_DOCUMENT_INTELLIGENCE_API_KEY = os.getenv("DOCUMENT_INTELLIGENCE_KEY")
GROQ_API = os.getenv("GROQ_API")

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text


# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(text)


# Function to call Azure Document Intelligence OCR
def call_document_intelligence(file_path):
    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT,
        credential=AzureKeyCredential(AZURE_DOCUMENT_INTELLIGENCE_API_KEY),
    )

    with open(file_path, "rb") as file:
        poller = document_intelligence_client.begin_analyze_document(
            model_id="prebuilt-read",
            body=file,
        )
    result = poller.result()

    if result:
        return result.content
    else:
        st.error("Document Intelligence processing failed.")
        return None


# Function to call Azure OpenAI GPT model
def call_azure_openai_gpt(content):
    llm = ChatGroq(api_key=GROQ_API,model="llama-3.3-70b-versatile",temperature=0.2)

    prompt = f"""
    
The following is the content extracted from a resume:
{content}

Extract the following details and return ONLY the key-value pairs in JSON format. Do not include any additional text or explanations:
- Name
- Years of Experience
- Linkedin Profile
- Contact Number
- Address
- Email
- Skills

    """

    prompt = PromptTemplate(
        input_variables=["content"],
        template=prompt,
    )
    chain = prompt | llm
    # Execute the chain
    try:
        result = chain.invoke({"content": content})
        return result.content
    except Exception as e:
        print("Error:", e)
        raise


# Streamlit app
st.title("Resume Processing App")

uploaded_file = st.file_uploader(
    "Attach Resume File", type=["pdf", "docx", "txt", "jpeg", "png"]
)
if st.button("Submit"):
    if uploaded_file is not None:
        with st.spinner("Processing your file..."):
            # Save uploaded file
            file_path = os.path.join("resumes", uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            # Determine file type and extract text
            file_extension = uploaded_file.name.split(".")[-1].lower()
            content = ""

            if file_extension == "pdf":
                content = extract_text_from_pdf(file_path)
            elif file_extension == "docx":
                content = extract_text_from_docx(file_path)
            elif file_extension == "txt":
                with open(file_path, "r", encoding="utf-8") as file:
                    content = file.read()
            else:  # Use OCR for other file types (e.g., images)
                content = call_document_intelligence(file_path)
                # st.write(content)
            if content:
                # Call Azure OpenAI to process content
                gpt_output = call_azure_openai_gpt(content)
                # st.write(gpt_output)
                match = re.search(r"\{.*\}", gpt_output, re.DOTALL)
                # Extract JSON from GPT output
                # try:
                    # gpt_output = str(gpt_output)
                    # st.write(type(gpt_output))
                if match:
                    extracted_data = json.loads(match.group(0))
                else:
                    st.error("Could not extract JSON from GPT output.")
                    # st.write("GPT Output:", gpt_output)
                # extracted_data = json.loads(gpt_output)
                # st.write(extracted_data)
                # Display fields
                st.subheader("Extracted Information")
                
                st.text_input("Name", value=extracted_data.get("Name", ""), key="name")
                years_of_experience_raw = extracted_data.get(
                    "Years of Experience", "0"
                )
                years_of_experience_cleaned = int(
                    "".join(filter(str.isdigit, years_of_experience_raw))
                )
                st.number_input(
                    "Years of Experience",
                    value=years_of_experience_cleaned,
                    key="experience",
                    min_value=0,
                )
                st.text_input(
                    "Contact Number",
                    value=extracted_data.get("Contact Number", ""),
                    key="contact_number",
                )
                st.text_input("Linkedin Profile", value=extracted_data.get("Linkedin Profile", ""), key="linkedin_profile")
                st.text_input("Email", value=extracted_data.get("Email", ""), key="email")
                st.text_area(
                    "Skills", value=extracted_data.get("Skills", ""), key="skills"
                )
                st.text_area(
                    "Address", value=extracted_data.get("Address", ""), key="Address"
                )
                    
                        
                # except json.JSONDecodeError:
                #     st.error("Failed to parse GPT output. Please check the response.")
            else:
                st.error("Failed to extract content from the uploaded file.")
    else:
        st.warning("Please upload a file before submitting.")






# from langchain.prompts import PromptTemplate
# from langchain.chains import LLMChain
# import os
# import streamlit as st
# from azure.core.credentials import AzureKeyCredential
# from azure.ai.documentintelligence import DocumentIntelligenceClient
# from langchain_groq import ChatGroq
# from openai import AzureOpenAI
# from PyPDF2 import PdfReader
# from docx import Document
# from dotenv import load_dotenv
# import json

# # Load environment variables
# load_dotenv()

# # Azure API configuration
# AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT = os.getenv("DOCUMENT_INTELLIGENCE_ENDPOINT")
# AZURE_DOCUMENT_INTELLIGENCE_API_KEY = os.getenv("DOCUMENT_INTELLIGENCE_KEY")
# AZURE_OPENAI_ENDPOINT = os.getenv("OPENAI_ENDPOINT")
# AZURE_OPENAI_API_KEY = os.getenv("OPENAI_KEY")
# AZURE_OPENAI_DEPLOYMENT_NAME = os.getenv("OPENAI_DEPLOYMENT_NAME")
# GROQ_API = os.getenv("GROQ_API")


# # Function to extract text from PDF
# def extract_text_from_pdf(file_path):
#     reader = PdfReader(file_path)
#     text = ""
#     for page in reader.pages:
#         text += page.extract_text()
#     return text


# # Function to extract text from DOCX
# def extract_text_from_docx(file_path):
#     doc = Document(file_path)
#     text = [p.text for p in doc.paragraphs if p.text.strip()]
#     return "\n".join(text)


# # Function to call Azure Document Intelligence OCR
# def call_document_intelligence(file_path):
#     document_intelligence_client = DocumentIntelligenceClient(
#         endpoint=AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT,
#         credential=AzureKeyCredential(AZURE_DOCUMENT_INTELLIGENCE_API_KEY),
#     )

#     with open(file_path, "rb") as file:
#         poller = document_intelligence_client.begin_analyze_document(
#             model_id="prebuilt-read",
#             document=file,
#         )
#     result = poller.result()

#     if result:
#         return result.content
#     else:
#         st.error("Document Intelligence processing failed.")
#         return None


# # Function to call Azure OpenAI GPT model
# def call_azure_openai_gpt(content):
#     # llm = ChatGroq(api_key=GROQ_API,model="llama-3.2-1b-preview",temperature=0.2)
#     client = AzureOpenAI(
#         azure_endpoint=AZURE_OPENAI_ENDPOINT,
#         api_key=AZURE_OPENAI_API_KEY,
#         api_version="2024-05-01-preview",
#     )
#     # # Define the prompt template
#     # template = """
#     # The following is the content extracted from a resume:
#     # {content}
#     # Please extract the following details and return the output as key-value pairs in JSON format and always keep the json hierarchy:
#     # - Name
#     # - Contact Number
#     # - Email
#     # - Address
#     # - LinkedIn Profile
#     # - Years of Experience
#     # - Education (up to 3 entries with degree name and institution name)
#     # - Experience Details (company name and title for each position)
#     # - Achievements
#     # - Certifications
#     # """
#     prompt = f"""
#     The following is the content extracted from a resume:
#     {content}

#     Please extract the following details and return them as key-value pairs in JSON format:
#     - Name
#     - Contact Number
#     - Email
#     - Address
#     - LinkedIn Profile
#     - Years of Experience
#     - Education (up to 3 entries with degree name and institution name)
#     - Experience Details (company name and title for each position)
#     - Achievements
#     - Certifications
#     """

#     chat_prompt = [
#         {
#             "role": "system",
#             "content": "You are an AI assistant that helps people extract structured data from resumes.",
#         },
#         {
#             "role": "user",
#             "content": prompt,
#         },
#     ]
#     # Adjust the format if `role` is not supported
#     # chat_prompt = [
#     #     {"type": "instruction", "content": "You are a resume parser."},
#     #     {"type": "input", "content": prompt}
#     # ]

#     completion = client.chat.completions.create(
#         model=AZURE_OPENAI_DEPLOYMENT_NAME,
#         messages=chat_prompt,
#         max_tokens=800,  
#         temperature=0.4,  
#         top_p=0.95,  
#         frequency_penalty=0,  
#         presence_penalty=0,
#         stop=None,  
#         stream=False
#     )
# # Generate completion using ChatGroq

#     # completion = llm.generate(
#     #     messages=chat_prompt,
#     #     max_tokens=1200,  # Set token limit
#     #     temperature=0.5,  # Temperature for randomness
#     #     top_p=0.95,  # Top-p sampling
#     # )

#     return completion.choices[0].message.content
#      # Create the PromptTemplate
#     # prompt = PromptTemplate(
#     #     input_variables=["content"],
#     #     template=template,
#     # )
    
#     # Create the LLMChain
#     # chain = LLMChain(llm=llm, prompt=prompt)
    
#     template = """
#     # The following is the content extracted from a resume:
#     # {content}
#     # Please extract the following details and return the output as key-value pairs in JSON format and always keep the json hierarchy:
#     # - Name
#     # - Contact Number
#     # - Email
#     # - Address
#     # - LinkedIn Profile
#     # - Years of Experience
#     # - Education (up to 3 entries with degree name and institution name)
#     # - Experience Details (company name and title for each position)
#     # - Achievements
#     # - Certifications
#     # """


# # Streamlit app
# st.title("Enhanced Resume Processing App")

# uploaded_file = st.file_uploader(
#     "Attach Resume File", type=["pdf", "docx", "txt", "jpeg", "png"]
# )
# if st.button("Submit"):
#     if uploaded_file is not None:
#         with st.spinner("Processing your file..."):
#             # Save uploaded file
#             file_path = os.path.join("resumes", uploaded_file.name)
#             with open(file_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())

#             # Determine file type and extract text
#             file_extension = uploaded_file.name.split(".")[-1].lower()
#             content = ""

#             if file_extension == "pdf":
#                 content = extract_text_from_pdf(file_path)
#             elif file_extension == "docx":
#                 content = extract_text_from_docx(file_path)
#             elif file_extension == "txt":
#                 with open(file_path, "r", encoding="utf-8") as file:
#                     content = file.read()
#             else:  # Use OCR for other file types (e.g., images)
#                 content = call_document_intelligence(file_path)

#             if content:
#                 # Call Azure OpenAI to process content
#                 gpt_output = call_azure_openai_gpt(content)
#                 st.write(gpt_output)
#                 # Extract JSON from GPT output
#                 try:
#                     extracted_data = json.loads(gpt_output)

#                     # Display fields
#                     st.subheader("Extracted Information")
#                     st.text_input("Name", value=extracted_data.get("Name", "N/A"))
#                     st.text_input(
#                         "Contact Number",
#                         value=extracted_data.get("Contact Number", "N/A"),
#                     )
#                     st.text_input("Email", value=extracted_data.get("Email", "N/A"))
#                     st.text_area("Address", value=extracted_data.get("Address", "N/A"))
#                     st.text_input(
#                         "LinkedIn Profile",
#                         value=extracted_data.get("LinkedIn Profile", "N/A"),
#                     )
#                     st.number_input(
#                         "Years of Experience",
#                         value=int(
#                             "".join(
#                                 filter(
#                                     str.isdigit,
#                                     extracted_data.get("Years of Experience", "0"),
#                                 )
#                             )
#                         ),
#                         min_value=0,
#                     )

#                     # Education details
#                     st.subheader("Education")
#                     education = extracted_data.get("Education", [])
#                     for i in range(3):
#                         degree = education[i].get("Degree", "N/A") if i < len(education) else "N/A"
#                         institution = education[i].get("Institution", "N/A") if i < len(education) else "N/A"
#                         st.text_input(f"Degree {i + 1}", value=degree)
#                         st.text_input(f"Institution {i + 1}", value=institution)

#                     # Experience details
#                     st.subheader("Experience")
#                     experience = extracted_data.get("Experience Details", [])
#                     for i, exp in enumerate(experience):
#                         st.text_input(f"Company {i + 1}", value=exp.get("Company", "N/A"))
#                         st.text_input(f"Title {i + 1}", value=exp.get("Title", "N/A"))

#                     # Achievements
#                     st.subheader("Achievements")
#                     achievements = extracted_data.get("Achievements", "N/A")
#                     st.text_area("Achievements", value=achievements)

#                     # Certifications
#                     st.subheader("Certifications")
#                     certifications = extracted_data.get("Certifications", "N/A")
#                     st.text_area("Certifications", value=certifications)
#                 except json.JSONDecodeError:
#                     st.error("Failed to parse GPT output. Please check the response.")
#             else:
#                 st.error("Failed to extract content from the uploaded file.")
#     else:
#         st.warning("Please upload a file before submitting.")
