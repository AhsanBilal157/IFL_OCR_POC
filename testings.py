# import os
# import streamlit as st
# from azure.core.credentials import AzureKeyCredential
# from azure.ai.documentintelligence import DocumentIntelligenceClient
# from openai import AzureOpenAI
# from PyPDF2 import PdfReader
# from docx import Document
# from dotenv import load_dotenv
# import json
# import nltk
# from nltk.corpus import stopwords
# import tiktoken

# # Load environment variables
# load_dotenv()
# nltk.download("stopwords")
# stop_words = set(stopwords.words("english"))

# def preprocess_text(text):
#     words = text.split()
#     filtered_text = " ".join(word for word in words if word.lower() not in stop_words)
#     return filtered_text

# def limit_tokens(text, max_tokens=950):
#     encoding = tiktoken.get_encoding("cl100k_base")  # Encoding for GPT models
#     tokens = encoding.encode(text)
#     return encoding.decode(tokens[:max_tokens])

# def extract_text_from_pdf(file_path):
#     reader = PdfReader(file_path)
#     text = reader.pages[0].extract_text() if reader.pages else ""
#     return preprocess_text(text)

# def extract_text_from_docx(file_path):
#     doc = Document(file_path)
#     text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
#     return preprocess_text(text.split("\n", 1)[0])  # First page only

# def extract_text_from_txt(file_path):
#     with open(file_path, "r", encoding="utf-8") as file:
#         text = file.read()
#     return preprocess_text(text.split("\n", 1)[0])  # First page only

# def call_document_intelligence(file_path):
#     document_intelligence_client = DocumentIntelligenceClient(
#         endpoint=os.getenv("DOCUMENT_INTELLIGENCE_ENDPOINT"),
#         credential=AzureKeyCredential(os.getenv("DOCUMENT_INTELLIGENCE_KEY")),
#     )
#     with open(file_path, "rb") as file:
#         poller = document_intelligence_client.begin_analyze_document(
#             model_id="prebuilt-read", body=file
#         )
#     result = poller.result()
#     return preprocess_text(result.content) if result else None

# def call_azure_openai_gpt(content):
#     client = AzureOpenAI(
#         azure_endpoint=os.getenv("OPENAI_ENDPOINT"),
#         api_key=os.getenv("OPENAI_KEY"),
#         api_version="2024-08-01-preview",
#     )
#     content = limit_tokens(content, 950)
#     chat_prompt = [
#         {"role": "system", "content": "You are an AI assistant extracting structured information from resumes."},
#         {"role": "user", "content": content},
#     ]
#     completion = client.chat.completions.create(
#         model=os.getenv("OPENAI_DEPLOYMENT_NAME"),
#         messages=chat_prompt,
#         max_tokens=800,
#     )
#     return completion.choices[0].message.content

# st.title("Resume Processing App")
# uploaded_file = st.file_uploader("Attach Resume File", type=["pdf", "docx", "txt", "jpeg", "png"])
# if st.button("Submit"):
#     if uploaded_file:
#         with st.spinner("Processing your file..."):
#             file_path = os.path.join("resumes", uploaded_file.name)
#             with open(file_path, "wb") as f:
#                 f.write(uploaded_file.getbuffer())
            
#             ext = uploaded_file.name.split(".")[-1].lower()
#             content = ""
#             if ext == "pdf":
#                 content = extract_text_from_pdf(file_path)
#             elif ext == "docx":
#                 content = extract_text_from_docx(file_path)
#             elif ext == "txt":
#                 content = extract_text_from_txt(file_path)
#             else:
#                 content = call_document_intelligence(file_path)

#             if content:
#                 gpt_output = call_azure_openai_gpt(content)
#                 try:
#                     extracted_data = json.loads(gpt_output)
#                     st.subheader("Extracted Information")
#                     st.text_input("Name", value=extracted_data.get("Name", ""))
#                     st.number_input("Years of Experience", value=int(extracted_data.get("Years of Experience", 0)), min_value=0)
#                     st.text_input("Contact Number", value=extracted_data.get("Contact Number", ""))
#                     st.text_input("Linkedin Profile", value=extracted_data.get("Linkedin Profile", ""))
#                     st.text_input("Email", value=extracted_data.get("Email", ""))
#                     st.text_area("Skills", value=extracted_data.get("Skills", ""))
#                     st.text_area("Address", value=extracted_data.get("Address", ""))
#                 except json.JSONDecodeError:
#                     st.error("Failed to parse GPT output. Please check the response.")
#             else:
#                 st.error("Failed to extract content from the uploaded file.")
#     else:
#         st.warning("Please upload a file before submitting.")

import os
import streamlit as st
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from PyPDF2 import PdfReader
from docx import Document
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain.chains.llm import LLMChain
from langchain_groq import ChatGroq
import json
import re
import math

# Load environment variables
load_dotenv()

# Azure API configuration
AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT = os.getenv("DOCUMENT_INTELLIGENCE_ENDPOINT")
AZURE_DOCUMENT_INTELLIGENCE_API_KEY = os.getenv("DOCUMENT_INTELLIGENCE_KEY")
GROQ_API = os.getenv("GROQ_API")

# Function to extract text from PDF
def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = "".join([page.extract_text() for page in reader.pages if page.extract_text()])
    return text

# Function to extract text from DOCX
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    print(text)
    return text

# Function to call Azure Document Intelligence OCR
def call_document_intelligence(file_path):
    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT,
        credential=AzureKeyCredential(AZURE_DOCUMENT_INTELLIGENCE_API_KEY),
    )
    with open(file_path, "rb") as file:
        poller = document_intelligence_client.begin_analyze_document(
            model_id="prebuilt-read", body=file
        )
    result = poller.result()
    return result.content if result else None

# Function to call Azure OpenAI GPT model
def call_azure_openai_gpt(content):
    llm = ChatGroq(api_key=GROQ_API, model="llama-3.3-70b-versatile", temperature=0.2)
    #(calculated based on the joining and leaving dates from the Experience Details section)
    prompt = f"""
    The following is the content extracted from a resume:
    {content}

    Please extract and return ONLY the key-value pairs in JSON format with the following fields:

    Name
    Years of Experience 
    Linkedin Profile
    Contact Number
    Address
    Email
    Skills (list)
    Education (list of dictionaries with "Degree" and "Institution")
    Experience Details (list of dictionaries with "Company", "Title", "Joining Date", and "Leaving Date")
    Achievements (list)
    Certifications (list)
    """
    
    prompt = PromptTemplate(input_variables=["content"], template=prompt)
    chain = prompt | llm
    try:
        result = chain.invoke({"content": content})
        return result.content
    except Exception as e:
        print("Error:", e)
        return None

# Streamlit app
st.title("Resume Processing App")

uploaded_file = st.file_uploader("Attach Resume File", type=["pdf", "docx", "txt", "jpeg", "png"])
if st.button("Submit"):
    if uploaded_file is not None:
        with st.spinner("Processing your file..."):
            file_path = os.path.join("resumes", uploaded_file.name)
            with open(file_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            file_extension = uploaded_file.name.split(".")[-1].lower()
            content = ""

            if file_extension == "pdf":
                content = extract_text_from_pdf(file_path)
            # elif file_extension == "docx":
            #     content = extract_text_from_docx(file_path)
            elif file_extension == "txt" :
                with open(file_path, "r", encoding="utf-8") as file:
                    content = file.read()
            else:
                content = call_document_intelligence(file_path)
            
            if content:
                gpt_output = call_azure_openai_gpt(content)
                match = re.search(r"\{.*\}", gpt_output, re.DOTALL)

                if match:
                    try:
                        extracted_data = json.loads(match.group(0))
                    except json.JSONDecodeError:
                        extracted_data = {}
                else:
                    st.error("Could not extract JSON from GPT output.")
                    extracted_data = {}
                
                st.subheader("Extracted Information")
                with st.form("resume_details"):
                    if "Name" in extracted_data:
                        st.text_input("Name", value=extracted_data.get("Name", "N/A"), key="name")
                    
                    if "Years of Experience" in extracted_data:
                        st.text_input("Years of Experience", value=extracted_data.get("Years of Experience", "N/A"), key="experience")
                    
                    if "Contact Number" in extracted_data:
                        st.text_input("Contact Number", value=extracted_data.get("Contact Number", "N/A"), key="contact_number")
                    
                    if "Linkedin Profile" in extracted_data:
                        st.text_input("Linkedin Profile", value=extracted_data.get("Linkedin Profile", "N/A"), key="linkedin_profile")
                    
                    if "Email" in extracted_data:
                        st.text_input("Email", value=extracted_data.get("Email", "N/A"), key="email")
                    
                    if "Address" in extracted_data:
                        st.text_input("Address", value=extracted_data.get("Address", "N/A"), key="address")
                    
                    for field in ["Skills", "Certifications", "Achievements"]:
                        if field in extracted_data and isinstance(extracted_data[field], list):
                            st.text_area(field, value=", ".join(extracted_data[field]), key=field.lower())
                    
                    for section, fields in {"Education": ["Degree", "Institution"], "Experience Details": ["Company", "Title"]}.items():
                        if section in extracted_data and isinstance(extracted_data[section], list):
                            st.subheader(section)
                            for i, entry in enumerate(extracted_data[section]):
                                for field in fields:
                                    value = entry.get(field, "N/A")
                                    st.text_input(f"{field} {i+1}", value=value, key=f"{field.lower()}_{i}")
                    
                    submitted = st.form_submit_button("Submit")
            else:
                st.error("Failed to extract content from the uploaded file.")
    else:
        st.warning("Please upload a file before submitting.")
